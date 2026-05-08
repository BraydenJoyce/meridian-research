"""Converts research report markdown to PDF bytes (ADR-007)."""
from __future__ import annotations

import io
import re

import mistune
from fpdf import FPDF


def generate_pdf(markdown: str) -> bytes:
    """Convert a markdown report string to PDF bytes.

    Uses mistune to parse markdown to HTML, then renders to PDF via fpdf2.
    Supports: ## and ### headings, bullet lists, and paragraph text.

    Args:
        markdown: Raw markdown report text.

    Returns:
        PDF file content as bytes (starts with b'%PDF').
    """
    html = mistune.html(markdown)
    lines = _html_to_lines(html)

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    for kind, text in lines:
        if kind == "h2":
            pdf.set_font("Helvetica", style="B", size=14)
            pdf.ln(4)
            pdf.multi_cell(0, 8, text)
            pdf.ln(2)
        elif kind == "h3":
            pdf.set_font("Helvetica", style="B", size=12)
            pdf.multi_cell(0, 7, text)
            pdf.ln(1)
        elif kind == "li":
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, f"  •  {text}")
        else:
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, text)
            pdf.ln(2)

    return bytes(pdf.output())


_TAG_RE = re.compile(r"<[^>]+>")
_LINK_RE = re.compile(r'href="([^"]+)"')


def _html_to_lines(html: str) -> list[tuple[str, str]]:
    lines: list[tuple[str, str]] = []
    pos = 0
    tag_re = re.compile(r"<(/?)(\w+)[^>]*>", re.IGNORECASE)

    for m in tag_re.finditer(html):
        before = html[pos : m.start()]
        if before.strip():
            lines.append(("p", _strip_tags(before).strip()))
        pos = m.end()

        closing = m.group(1) == "/"
        tag = m.group(2).lower()

        if not closing:
            if tag in ("h1", "h2"):
                end = html.find(f"</{tag}>", pos)
                if end != -1:
                    content = _strip_tags(html[pos:end]).strip()
                    lines.append(("h2", content))
                    pos = end + len(f"</{tag}>")
            elif tag == "h3":
                end = html.find("</h3>", pos)
                if end != -1:
                    content = _strip_tags(html[pos:end]).strip()
                    lines.append(("h3", content))
                    pos = end + 5
            elif tag == "li":
                end = html.find("</li>", pos)
                if end != -1:
                    content = _strip_tags(html[pos:end]).strip()
                    lines.append(("li", content))
                    pos = end + 5

    remainder = html[pos:]
    if remainder.strip():
        lines.append(("p", _strip_tags(remainder).strip()))

    return [(k, t) for k, t in lines if t]


def _strip_tags(text: str) -> str:
    return _TAG_RE.sub("", text)


def to_docx(markdown: str, title: str = "Intelligence Report") -> bytes:
    """Convert a markdown report string to DOCX bytes.

    Uses mistune to parse markdown to HTML, then renders to a Word document
    via python-docx.  Supports: ## and ### headings, bullet lists, and
    paragraph text.

    Args:
        markdown: Raw markdown report text.
        title:    Document title added as a top-level heading.

    Returns:
        DOCX file content as bytes.
    """
    from docx import Document  # noqa: PLC0415  (lazy import – optional dep)

    doc = Document()
    doc.add_heading(title, level=0)

    html = mistune.html(markdown)
    for kind, text in _html_to_lines(html):
        if kind == "h2":
            doc.add_heading(text, level=1)
        elif kind == "h3":
            doc.add_heading(text, level=2)
        elif kind == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
